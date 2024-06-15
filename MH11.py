from docx import Document
from copy import deepcopy
import os
import re
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, parse_xml
# Function to remove hyperlinks
def remove_hyperlinks(paragraph):
    for run in paragraph.runs:
        r_element = run._element
        hyperlink = r_element.xpath(".//w:hyperlink")
        for link in hyperlink:
            link.getparent().remove(link)

# Load the Word document

def replacetext(transformer_market,company, doc):
    toc1count = 0
    toc1active = 0
    toc2count = 0
    firsttoc2 = 0
    current_text = None
    removetoc2 = None
    toc1active = False
    toc2active = False
    current_toc1 = None
    start_changing = False
    current_toc2 = None
    changes = []  # Track changes for global replacement
    remove_extra_tocs = []
    oldtoc1 = []
    company_changes = []
    removed_toc1 = []
    company_active = False
    for paragraph in doc.paragraphs:
        if paragraph.style.name == "toc 1":
            start_changing = True
        if start_changing:
            if toc1count < len(transformer_market):
                if paragraph.style.name == "toc 1" and toc1count==10:#this is  last toc1count by to replace

                    print("count6",toc2count," ",len(transformer_market[toc1count-1]))
                    if toc2count>1 and toc2count<(len(transformer_market[toc1count-1])-1):
                        toc2count+=1
                        for c in range(toc2count,len(transformer_market[toc1count-1])):
                            addtoc2(paragraph,current_toc2,toc1count,toc2count)
                            toc2count+=1
                    toc1active = False
                    toc2active = False
                if paragraph.style.name == "toc 1" and toc1count<10:

                    print("count",toc2count," ",len(transformer_market[toc1count-1]))
                    if toc2count>1 and toc2count<(len(transformer_market[toc1count-1])-1):
                        toc2count+=1
                        for c in range(toc2count,len(transformer_market[toc1count-1])):
                            addtoc2(paragraph,current_toc2,toc1count,toc2count)
                            toc2count+=1
                    toc1active = False
                    toc2active = False
                    if 'by' in paragraph.text.lower():

                        toc1active = True
                        firsttoc2 = 0
                        toc2count = 0
                        match = paragraph.text.strip()
                        print("matching", match)
                        pattern = re.compile(r'ANALYSIS,\sBy\s([^\t\n]+)')

                        # Find the match
                        match = pattern.search(paragraph.text)
                        if match:
                            print(match)
                            current_text = match[0].split('By ')[1]
                            paragraph.text = paragraph.text.replace(current_text, transformer_market[toc1count][0])
                            print("Checking TOC entry:", current_text, ", replace:", paragraph.text)  # Debugging output
                            changes.append(("By "+ current_text, "By "+transformer_market[toc1count][0]))
                            oldtoc1.append(current_text)
                            current_toc1 = paragraph

                            modified = True
                            toc1count+=1
                            print("toc1", toc1count)
                elif toc1count>9 and paragraph.style.name == "toc 1": #toc1count is the no of by replace-1
                    print("adding")
                    initial_toc1 = toc1count
                    for r in range(initial_toc1,len(transformer_market)):
                        addtoc1(paragraph, current_toc1,toc1count)
                        toc1count+=1
                        toc2count = 1
                        for c in range(toc2count, len(transformer_market[r])):
                            addtoc2(paragraph,current_toc2,toc1count,toc2count)
                            toc2count+=1
                if paragraph.style.name == "toc 2":
                    if toc1active:
                        firsttoc2 += 1
                        if 'global' in paragraph.text.lower().strip():

                            pattern = re.compile(r'Analysis,\sBY\s([^\t\n]+)')
                            # Find the match
                            match = pattern.search(paragraph.text)
                            print(match,":match", paragraph.text)
                            if match:
                                global_replace = match[0].split('BY ')[1]
                                paragraph.text = paragraph.text.replace(global_replace, transformer_market[toc1count-1][0])
                                print("global replace", global_replace, ", replace:", paragraph.text)
                        print("remove", toc2count,len(transformer_market[toc1count-1]))
                        if toc2count>len(transformer_market[toc1count-1])-2:
                            p_element = paragraph._element
                            p_element.getparent().remove(p_element)
                            continue
                        if firsttoc2 > 2:
                            print("Toc1 count:", toc1count-1, "Toc2 count:", toc2count+1)
                            new_text = transformer_market[toc1count-1][toc2count+1]
                            match = paragraph.text.strip()

                            match = re.match(r'(\d*.\d+)\s+(.*?)\s(\d+\s*$)', match)
                            print("regexmatching", match)
                            if match:
                                current_text = match[2].strip()
                                print("regex", current_text)
                            match = None
                            print(f"Modifying TOc2 entry from '{paragraph.text}' to '{new_text}'")
                            paragraph.text = paragraph.text.replace(current_text, new_text)
                            print(f"Modifying TOc2 entry from '{paragraph.text}")
                            changes.append((current_text, new_text))
                            current_toc2 = paragraph

                            # if match:
                            #     current_text = match[0].strip()
                            #     print("regextext", match[0])
                            toc2count+=1
            elif paragraph.style.name == "toc 1" and toc1count == len(transformer_market):
                removetoc2 = False

                pattern = re.compile(r'ANALYSIS,\sBy\s([^\t\n]+)')
                # Find the match
                match = pattern.search(paragraph.text)
                if match:
                    p_element = paragraph._element
                    remove_extra_tocs.append(match[0].split("By ")[1])
                    p_element.getparent().remove(p_element)
                    removetoc2 = True
                    continue
                else:
                    toc1count+=1
            if paragraph.style.name == "toc 2" and toc1count == len(transformer_market) :
                print("executive", paragraph.text)
                if removetoc2:
                    p_element = paragraph._element
                    if (not 'overview' in paragraph.text.lower().strip()) and (not 'global' in paragraph.text.lower().strip()):
                        match = re.match(r'(\d*.\d+)\s+(.*?)\s(\d+\s*$)', paragraph.text.strip())
                        remove_text = None
                        if match:
                            remove_text = match[2].strip()
                            print("regex", remove_text)
                            remove_extra_tocs.append(remove_text)

                    p_element.getparent().remove(p_element)
                    continue
                if toc1active:
                    firsttoc2 += 1
                    if 'global' in paragraph.text.lower().strip():

                        pattern = re.compile(r'Analysis,\sBY\s([^\t\n]+)')
                        # Find the match
                        match = pattern.search(paragraph.text)
                        print(match, ":match", paragraph.text)
                        if match:
                            global_replace = match[0].split('BY ')[1]
                            paragraph.text = paragraph.text.replace(global_replace, transformer_market[toc1count - 1][0])
                            print("global replace", global_replace, ", replace:", paragraph.text)
                    print("remove", toc2count,len(transformer_market[toc1count-1]))
                    if toc2count>len(transformer_market[toc1count-1])-2:
                        p_element = paragraph._element
                        p_element.getparent().remove(p_element)
                        continue
                    if firsttoc2 > 2:
                        print("Toc1 count:", toc1count-1, "Toc2 count:", toc2count+1)
                        new_text = transformer_market[toc1count-1][toc2count+1]
                        match = paragraph.text.strip()

                        match = re.match(r'(\d*.\d+)\s+(.*?)\s(\d+\s*$)', match)
                        print("regexmatching", match)
                        if match:
                            current_text = match[2].strip()
                            print("regex", current_text)
                        match = None
                        print(f"Modifying TOc2 entry from '{paragraph.text}' to '{new_text}'")
                        paragraph.text = paragraph.text.replace(current_text, new_text)
                        print(f"Modifying TOc2 entry from '{paragraph.text}")
                        changes.append((current_text, new_text))
                        current_toc2 = paragraph

                        # if match:
                        #     current_text = match[0].strip()
                        #     print("regextext", match[0])
                        toc2count+=1
            if paragraph.style.name == "toc 2" and toc1count> len(transformer_market) :
                print("executive", paragraph.text)
                for i, c in changes:
                    print("Company figure", paragraph.text, ":", i)
                    old_text = i
                    new_text = c

                    paragraph.text = paragraph.text.replace('By SEGMENT1', new_text)
                    print(paragraph.style.name, ":updated", paragraph.text)
                    break
            if 'toc 1' in paragraph.style.name and "company" in paragraph.text.lower().strip():
                company_active = True
                toc2count = 0
            if company_active and 'toc 2' in paragraph.style.name:
                if toc2count > len(company[0]) - 1:
                    p_element = paragraph._element
                    p_element.getparent().remove(p_element)
                    continue
                new_text = company[0][toc2count]
                match = paragraph.text.strip()

                match = re.match(r'(\d*.\d+)\s+(.*?)\s(\d+\s*$)', match)
                print("regexmatching", match)
                if match:
                    current_text = match[2].strip()
                    print("regex", current_text)
                match = None
                print(f"Modifying company entry from '{paragraph.text}' to '{new_text}'")
                paragraph.text = paragraph.text.replace(current_text, new_text)
                print(f"Modifying company entry from '{paragraph.text}")
                company_changes.append((current_text, new_text))
                toc2count+=1

        # if not paragraph.style.name in ["toc 1","toc 2"] and 'by' in :
    print(changes)
    print(remove_extra_tocs)
    print(company_changes)
    toc1count = 0
    toc2count = 0
    company_active = False
    last_run_properties_bolt = None
    bullettoc2count = 0
    webtoc2count = 0
    current_bullet_toc2 = None
    start_changing = False
    remove_extra = False
    remove_geography = False
    for i,paragraph in enumerate(doc.paragraphs):
        if paragraph.style.name == "toc 1":
            start_changing = True
        if start_changing:
            # print(paragraph.text,"sssss", paragraph.style.name)
            print("main count", toc1count)
            if 'main heading' in paragraph.style.name.lower() and toc1count==len(transformer_market):
                remove_extra = False
                print("main heading", paragraph.text, "remove main ", toc1count)
                pattern = re.compile(r'ANALYSIS,\sBy\s([^\t\n]+)')

                # Find the match
                match = pattern.search(paragraph.text)
                if match:
                    p_element = paragraph._element
                    parent_element = p_element.getparent()
                    parent_element.remove(p_element)
                    remove_extra = True
                    continue
            if remove_extra:
                print("parent", paragraph.text)
                p_element = paragraph._element
                next_element = p_element.getnext()
                parent_element = p_element.getparent()
                try:
                    parent_element.remove(p_element)
                except:
                    continue
                drawing_elements = None
                if next_element is not None:
                    drawing_elements = next_element.findall('.//w:drawing', namespaces=next_element.nsmap)
                if next_element is not None and drawing_elements:
                    print("image remove")
                    parent_element.remove(next_element)
                continue
            if 'main heading' in paragraph.style.name.lower():
                if toc1count==10:#this is  last toc1count by to replace

                    print("count6",toc2count," ",len(transformer_market[toc1count-1]))
                    if toc2count>1 and toc2count<(len(transformer_market[toc1count-1])-1):
                        toc2count+=1
                        for c in range(toc2count,len(transformer_market[toc1count-1])):
                            new_toc2 = current_toc2.insert_paragraph_before(transformer_market[toc1count-1][toc2count], style=current_toc2.style)
                            tem_text = current_toc2.text
                            current_toc2.text = new_toc2.text
                            new_toc2.text = tem_text
                            toc2count+=1
                    if bullettoc2count>1 and bullettoc2count<(len(transformer_market[toc1count-1])-1):
                        bullettoc2count+=1
                        for c in range(bullettoc2count,len(transformer_market[toc1count-1])):
                            # last_run_properties = get_run_properties(current_bullet_toc2.runs[0])
                            print("last run",last_run_properties_bolt)
                            new_toc2 = current_bullet_toc2.insert_paragraph_before(transformer_market[toc1count-1][bullettoc2count], style=current_bullet_toc2.style)
                            tem_text = current_bullet_toc2.text
                            current_bullet_toc2.text = new_toc2.text
                            new_toc2.text = tem_text
                            print("last run", last_run_properties_bolt, new_toc2.text,":",current_bullet_toc2.text)
                            for run in current_bullet_toc2.runs:
                                # Copy font properties
                                apply_run_properties(run, last_run_properties_bolt)
                            for run in new_toc2.runs:
                                # Copy font properties
                                apply_run_properties(run, last_run_properties_bolt)
                            add_bullet_paragraph(new_toc2,current_bullet_toc2)
                            bullettoc2count+=1
                    toc1active = False
                    toc2active = False
                if toc1count<10:

                    print("count",toc2count," ",len(transformer_market[toc1count-1]))
                    if toc2count>1 and toc2count<(len(transformer_market[toc1count-1])-1):
                        toc2count+=1
                        for c in range(toc2count,len(transformer_market[toc1count-1])):
                            new_toc2 = current_toc2.insert_paragraph_before(transformer_market[toc1count - 1][toc2count],
                                                                            style=current_toc2.style)
                            tem_text = current_toc2.text
                            current_toc2.text = new_toc2.text
                            new_toc2.text = tem_text
                            toc2count+=1
                    if bullettoc2count>1 and bullettoc2count<(len(transformer_market[toc1count-1])-1):
                        bullettoc2count+=1
                        for c in range(bullettoc2count,len(transformer_market[toc1count-1])):
                            # last_run_properties = get_run_properties(current_bullet_toc2.runs[0])

                            new_toc2 = current_bullet_toc2.insert_paragraph_before(transformer_market[toc1count - 1][bullettoc2count]
                                                                            ,style = current_bullet_toc2.style)
                            tem_text = current_bullet_toc2.text
                            current_bullet_toc2.text = new_toc2.text

                            new_toc2.text = tem_text
                            print("last run", last_run_properties_bolt, new_toc2.text,":",current_bullet_toc2.text)
                            for run in current_bullet_toc2.runs:
                                # Copy font properties
                                apply_run_properties(run, last_run_properties_bolt)
                            for run in new_toc2.runs:
                                # Copy font properties
                                apply_run_properties(run, last_run_properties_bolt)
                            add_bullet_paragraph(new_toc2,current_bullet_toc2)
                            bullettoc2count+=1
                    toc1active = False
                    toc1active = False
                    if 'by' in paragraph.text.lower():
                        print("main heading", paragraph.text)
                        pattern = re.compile(r'ANALYSIS,\sBy\s([^\t\n]+)')

                        # Find the match
                        match = pattern.search(paragraph.text)
                        to_match = None
                        for old_text, new_text in changes:
                            if match:
                                to_match = match[0].split(", ")[1]
                                print("enter", to_match, ":", old_text)
                                if old_text.lower().strip() == to_match.lower().strip():
                                    # print(paragraph.text)
                                    print("enter1")
                                    current_toc1 = paragraph
                                    paragraph.text = paragraph.text.replace(current_toc1.text.split('By')[1], new_text.replace('By',''))
                                    toc1active = True
                                    toc2count = 0
                                    bullettoc2count = 0
                                    toc1count+=1
                                    break
                elif toc1count>9: #toc1count is the no of by replace-1
                    print("adding")
                    initial_toc1 = toc1count
                    for r in range(initial_toc1,len(transformer_market)):
                        new_toc1 = paragraph.insert_paragraph_before(current_toc1.text,style=current_toc1.style)
                        new_toc1.text = transformer_market[toc1count][0]
                        current_toc1 = new_toc1
                        toc1count+=1
                        toc2count = 1
                        for c in range(toc2count, len(transformer_market[r])):
                            new_toc2 = paragraph.insert_paragraph_before(current_toc2.text, style=current_toc2.style)
                            new_toc2.text = transformer_market[toc1count-1][toc2count]
                            current_toc2 = new_toc2
                            toc2count+=1
            if ('head 1' in paragraph.style.name.lower()) and toc1active and (not 'global' in paragraph.text.lower() )and (not'overview' in paragraph.text.lower()):
                print("remove", toc2count, len(transformer_market[toc1count - 1]))
                if toc2count > len(transformer_market[toc1count - 1]) - 2:
                    p_element = paragraph._element
                    next_element = p_element.getnext()
                    parent_element = p_element.getparent()
                    parent_element.remove(p_element)
                    drawing_elements = None
                    if next_element is not None:
                        drawing_elements = next_element.findall('.//w:drawing', namespaces=next_element.nsmap)
                    if next_element is not None and drawing_elements:
                        print("image remove")
                        parent_element.remove(next_element)
                    continue
                if toc2count<len(transformer_market[toc1count-1])-1:
                    print("head1", transformer_market[toc1count-1][toc2count+1])
                    paragraph.text = transformer_market[toc1count-1][toc2count+1]
                    current_toc2 = paragraph
                    toc2count+=1
            if ('list paragraph' in paragraph.style.name.lower()) and toc1active and (not 'global' in paragraph.text.lower() )and (not'overview' in paragraph.text.lower()):
                print("remove", bullettoc2count, len(transformer_market[toc1count - 1]))
                if bullettoc2count > len(transformer_market[toc1count - 1]) - 2:
                    p_element = paragraph._element
                    p_element.getparent().remove(p_element)
                    continue
                if bullettoc2count<len(transformer_market[toc1count-1])-1:
                    last_run_properties_bolt = get_run_properties(paragraph.runs[0])
                    print("list paragraph", transformer_market[toc1count-1][bullettoc2count+1])
                    paragraph.text = transformer_market[toc1count-1][bullettoc2count+1]
                    for run in paragraph.runs:
                        # Copy font properties
                        apply_run_properties(run, last_run_properties_bolt)
                    current_bullet_toc2 = paragraph
                    bullettoc2count+=1
            if "normal (web)" in paragraph.style.name.strip().lower():
                is_break = True
                if 'mainly split into segments' in paragraph.text.lower().strip():
                    is_break = False
                print(is_break, "web", paragraph.text)
                webtoc2count = 0
                for e,c in changes:
                    if webtoc2count == 2 and is_break:
                            break
                    if 'by' in e.strip().lower():
                        old_text = e.replace('By ','')
                        new_text = c.replace('By ','')
                    else:
                        old_text = e
                        new_text = c
                    pattern = re.compile(
                        r"(\s*)\b" + r"\s*".join(re.escape(word) for word in old_text.split()) + r"\b(\s*)",
                                             re.IGNORECASE)
                    match = re.search(pattern, paragraph.text)
                    if match:
                         match.group(0)
                         old_style = paragraph.style
                         last_run_properties = get_run_properties(paragraph.runs[0])
                         paragraph.text = paragraph.text.replace(match.group(0).strip(),new_text,1)
                         for run in paragraph.runs:
                             # Copy font properties
                              apply_run_properties(run, last_run_properties)

                         print("updated", paragraph.text, ":", match.group(0))
                         webtoc2count+=1
            if (("table_title" in paragraph.style.name.strip().lower()) or ("figure _ title" in paragraph.style.name.strip().lower()) or
                    ("head 1" in paragraph.style.name.strip().lower() and 'global' in paragraph.text.lower())):
                print(paragraph.style.name, ":", paragraph.text)
                for i in remove_extra_tocs:
                    if i.lower().strip() in paragraph.text.lower().strip():
                        p_element = paragraph._element
                        try:
                            p_element.getparent().remove(p_element)
                        except:
                            continue
                webtoc2count = 0
                for i,c in changes:
                    if webtoc2count == 2:
                        break
                    if 'by' in i.strip().lower():
                        old_text = i.replace('By','')
                        new_text = c.replace('By','')
                    else:
                        old_text = i
                        new_text = c
                    pattern = re.compile(
                        r"(\s*)\b" + r"\s*".join(re.escape(word) for word in old_text.split()) + r"\b(\s*)",
                                             re.IGNORECASE)
                    match = re.search(pattern, paragraph.text)
                    if match:
                        paragraph.text = paragraph.text.replace(match.group(0),new_text)
                        print(paragraph.style.name, ":updated", paragraph.text)
            if ("table of figures" in paragraph.style.name.strip().lower()):
                print(paragraph.style.name, ":", paragraph.text)
                for i in remove_extra_tocs:
                    if i.lower().strip() in paragraph.text.lower().strip():
                        p_element = paragraph._element
                        try:
                            p_element.getparent().remove(p_element)
                        except:
                            continue
                webtoc2count = 0
                for i, c in changes:
                    if webtoc2count == 2:
                        break
                    if 'by' in i.strip().lower():
                        old_text = i.replace('By', '')
                        new_text = c.replace('By', '')
                    else:
                        old_text = i
                        new_text = c
                    pattern = re.compile(
                        r"(\s*)\b" + r"\s*".join(re.escape(word) for word in old_text.split()) + r"\b(\s*)",
                        re.IGNORECASE)
                    match = re.search(pattern, paragraph.text)
                    if match:
                        if 'figure' in paragraph.text.lower():
                            paragraph.text = paragraph.text.replace(match.group(0), new_text + ' ')
                        else:
                            paragraph.text = paragraph.text.replace(match.group(0), new_text)
                        print(paragraph.style.name, ":updated", paragraph.text)
            if 'main heading' in paragraph.style.name.lower() and 'geography' in paragraph.text.lower().strip():
                print("geography True")
                remove_geography = True
            if remove_geography:
                print("geography", paragraph.text)
                for i in remove_extra_tocs:
                    if i.lower().strip() in paragraph.text.lower().strip():
                        p_element = paragraph._element
                        try:
                            p_element.getparent().remove(p_element)
                        except:
                            continue
                        continue
            if 'main heading' in paragraph.style.name.lower().strip() and "company" in paragraph.text.lower().strip():
                print("Company Active")
                remove_geography = False
                company_active = True
                toc2count = 0
            if company_active:
                print("company table", paragraph.style.name)
                if 'head 1' in paragraph.style.name.lower().strip():
                    print("remove", toc2count, len(company[0]))
                    if toc2count > len(company[0]) - 1:
                        p_element = paragraph._element
                        next_element = p_element.getnext()
                        parent_element = p_element.getparent()
                        parent_element.remove(p_element)
                        drawing_elements = None
                        if next_element is not None:
                            drawing_elements = next_element.findall('.//w:drawing', namespaces=next_element.nsmap)
                        if next_element is not None and drawing_elements:
                            print("image remove")
                            parent_element.remove(next_element)
                        continue
                    if toc2count < len(company[0]):
                        print("head1", company[0][toc2count])
                        paragraph.text = company[0][toc2count]
                        current_toc2 = paragraph
                        toc2count += 1
                if ("table_title" in paragraph.style.name.strip().lower()):
                    print("company table", paragraph.style.name)
                    for i, c in company_changes:
                        print("Company figure", paragraph.text, ":", i)
                        old_text = i
                        new_text = c

                        paragraph.text = paragraph.text.replace(old_text, new_text)
                        print(paragraph.style.name, ":updated", paragraph.text)
                        break
            if "table of figures" in paragraph.style.name.strip().lower() and ' roche ' in paragraph.text.lower().strip():
                print("company table", paragraph.style.name)
                for i, c in company_changes:
                    print("Company figure", paragraph.text, ":", i)
                    old_text = i
                    new_text = c

                    paragraph.text = paragraph.text.replace(old_text,new_text)
                    print(paragraph.style.name, ":updated", paragraph.text)
                    break


    table_remove = []
    for table_index, table in enumerate(doc.tables):
        exit_row = False
        move_newtable = False
        # Check the first cell of the first row
        first_cell_text = table.rows[0].cells[0].text.strip()
        for i in remove_extra_tocs:
            print("table remove", first_cell_text, ":", i)
            if i.lower().strip() in first_cell_text.lower():
                print("table remove")
                table._element.getparent().remove(table._element)
                move_newtable = True
                break
        if move_newtable:
            continue
        index = tablehelper(oldtoc1,first_cell_text)
        print(first_cell_text,f":Table {table_index + 1} matches the condition. Iterating over rows...",index)
        if not index==None:
            toc2count = 0

            # Iterate over all rows in the current table, excluding the last row
            for row_idx, row in enumerate(table.rows[:-1]):

                if exit_row:
                    break
                # Get the first cell of the current row
                cell = row.cells[0]
                # Iterate over all paragraphs in the first cell
                for paragraph_idx, paragraph in enumerate(cell.paragraphs):
                     if row_idx==0:

                        # print(paragraph.text,":",transformer_market[index][0])
                        replace_text_in_paragraph(paragraph, paragraph.text, transformer_market[index][0])
                        continue
                        # print("toc1 replaced",paragraph.text)
                     else:
                        if toc2count > len(transformer_market[index]) - 2:
                            delete_counter = row_idx
                            # Check if the specified row index is within the table's bounds
                            if delete_counter < len(table.rows)-1:
                                print(delete_counter,len(table.rows))
                                # Collect rows to delete
                                rows_to_delete = []
                                for r in range(delete_counter, len(table.rows) - 1):
                                    rows_to_delete.append(table.rows[r])

                                # Delete the collected rows
                                for row in rows_to_delete:
                                    print("delete", row.cells[0].text)
                                    table._element.remove(row._element)

                                exit_row = True
                                break
                            else:
                                print(f"No row at index to delete.")
                                break  # Exit the loop
                        print(paragraph.text, ":", transformer_market[index][toc2count+1])
                        last_run_properties = get_run_properties(paragraph.runs[0])
                        paragraph.text = paragraph.text.replace(paragraph.text,
                                                                transformer_market[index][toc2count + 1])
                        for run in paragraph.runs:
                            # # Copy font properties
                            apply_run_properties(run, last_run_properties)
                        print("replace", paragraph.text)
                        toc2count += 1
                        if row_idx == len(table.rows) - 2:
                            for remaining_text in transformer_market[index][toc2count + 1:]:
                                add_row_after(table, row, remaining_text)
                                toc2count += 1
                            exit_row = True
                            break  # Exit the row loop to move to the next table

        else:
            print(f"Table {table_index + 1} does not match the condition. Skipping...")

    print(oldtoc1)


def get_run_properties(run):
    properties = {
        'bold': run.bold,
        'italic': run.italic,
        'underline': run.underline,
        'font_name': run.font.name,
        'font_size': run.font.size,
        'font_color': run.font.color.rgb if run.font.color else None,
    }
    return properties

def apply_run_properties(run, properties):
    run.bold = properties['bold']
    run.italic = properties['italic']
    run.underline = properties['underline']
    run.font.name = properties['font_name']
    run.font.size = properties['font_size']
    if properties['font_color']:
        run.font.color.rgb = properties['font_color']


def insert_text_with_format(paragraph, text):
    if not text:
        return
    run = paragraph.add_run()
    run.text = text


def add_bullet_paragraph(p,paragraph):
    last_pPr = paragraph._p.get_or_add_pPr()
    copied_pPr = copy_xml_element(last_pPr)
    new_pPr = p._p.get_or_add_pPr()
    new_pPr.append(copied_pPr)

def copy_xml_element(element):
    return parse_xml(element.xml)

def add_row_after(table, last_row, text):
    new_row = table.add_row()  # Add a new row at the end of the table
    print("Row adding")


    # Insert the new row after the specified row
    table._tbl.insert(len(table.rows), new_row._tr)


    for i, cell in enumerate(last_row.cells):
        # Remove the initial empty paragraph in the new row cell

        new_cell = new_row.cells[i]
        new_cell._element.clear_content()

        # Copy content and properties from each paragraph in the cell
        for paragraph in cell.paragraphs:
            new_paragraph = new_row.cells[i].add_paragraph()
            for run in paragraph.runs:
                new_run = new_paragraph.add_run(run.text)
                # Copy font properties
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                new_run.font.size = run.font.size
                new_run.font.name = run.font.name
                new_run.font.color.rgb = run.font.color.rgb

            # Copy paragraph style
            new_paragraph.style = paragraph.style

            # Copy paragraph alignment
            new_paragraph.alignment = paragraph.alignment
    replace_text_in_paragraph(new_row.cells[0].paragraphs[0], new_row.cells[0].paragraphs[0].text, text)

def replace_text_in_paragraph(paragraph, old_text, new_text):
    for run in paragraph.runs:
        if old_text.lower().strip() in run.text.lower().strip():
            run.text = run.text.replace(old_text.strip(), new_text.strip())
def replace_text_in_bullets(paragraph, old_text, new_text):
    for run in paragraph.runs:
        if old_text.lower().strip() in run.text.lower().strip():
            run.text = run.text.replace(old_text.strip(), new_text.strip())
def addtoc1(paragraph, current_toc1, toc1count):
    new_toc1 = paragraph.insert_paragraph_before(current_toc1.text,style=current_toc1.style)
    print("newtoc1", new_toc1.text)
    pattern = re.compile(r'^(\d+)\t(.+?)\t(\d+)$')

    # Find the matches
    match = pattern.match(current_toc1.text)
    # print("match", match)
    if match:
        current_text = match[2]
        print(current_text)
        new_toc1.text = new_toc1.text.replace(current_text, transformer_market[toc1count][0],1)
        new_toc1.text = new_toc1.text.replace(match[1], str(int(match[1])+1),1)
        print('newpara', new_toc1.text)
        current_toc1 = new_toc1
        # toc1count += 1
def tablehelper(oldtoc1, text):
    for i,toc1 in enumerate(oldtoc1):
        print("toc1",toc1.lower(),":",text.lower())
        if text.lower() in toc1.lower():
            print("true")
            return i

    return None
def addtoc2(paragraph, current_toc2, toc1count,toc2count):
    new_toc2 = paragraph.insert_paragraph_before(current_toc2.text, style=current_toc2.style)
    match = re.match(r'(\d*\.\d*)(.*)(\d{2})', current_toc2.text)
    # print("match",match)
    if match:
        current_text = match[2]
        print('digit', current_text)
        value = float(match[1]) + 0.1
        rounded_value = round(value, 1)
        result = str(rounded_value)
        new_toc2.text = new_toc2.text.replace(current_text, '\t'+transformer_market[toc1count-1][toc2count]+'\t',1)
        new_toc2.text = new_toc2.text.replace(match[1], result,1)
        print('newpara2', new_toc2.text)
        current_toc2 = new_toc2
        # toc2count += 1
def delete_segments_and_blank_pages(doc):
    for para in doc.paragraphs:
        if any(f"segment {i}" in para.text.lower() for i in range(1, 11)):
            last_run_properties_bolt = get_run_properties(para.runs[0])
            print("Original paragraph:", para.text, "Style:", para.style.name)
            new_text = re.sub(r',\s*Segment\s*\d+', '', para.text)
            para.text = new_text
            for run in para.runs:
                # Copy font properties
                apply_run_properties(run, last_run_properties_bolt)
            print("Modified paragraph:", para.text, "Style:", para.style.name)

    geography = False
    for i, paragraph in enumerate(doc.paragraphs):
        if 'main heading' in paragraph.style.name.lower().strip() and (
                "geography" in paragraph.text.lower().strip() or "market analysis" in paragraph.text.lower().strip()):
            print("geo", paragraph.text)
            geography = True
        if 'main heading' in paragraph.style.name.lower().strip() and (
                "landscape" in paragraph.text.lower().strip() or "executive summary" in paragraph.text.lower().strip()):
            print("geo", paragraph.text)
            geography = False
        if paragraph.style.name == 'Normal' and not paragraph.text.strip() and geography:
            try:

                p_element = paragraph._element
                drawing_elements = p_element.findall('.//w:drawing', namespaces=p_element.nsmap)
                if drawing_elements:
                    continue
                p_element.getparent().remove(p_element)
            except:
                pass

def get_segments():
    while True:
        try:
            num_segments = int(input("Enter the number of segments: "))
            if num_segments < 1:
                print("Number of segments must be at least 1. Please try again.")
                continue
            break
        except ValueError:
            print("Invalid input. Please enter a valid number.")

    segments = []

    for i in range(num_segments):
        while True:
            segment_input = input(
                f"Enter the segment name and its subsegments separated by commas, enclosed in double quotes (e.g., \"Offering\",\"Software\",\"Services\") for segment {i + 1}: ")
            segment_details = re.findall(r'"(.*?)"', segment_input)
            if len(segment_details) < 2:
                print("Invalid input. Each segment must have at least one subsegment. Please try again.")
            else:
                segments.append(segment_details)
                break

    return segments


def get_companies():
    while True:
        company_input = input(
            "Enter the company names separated by commas, enclosed in double quotes (e.g., \"Manhattan Associates\", \"Blue Yonder Group, Inc.\", \"SAP\"): ")
        companies = re.findall(r'"(.*?)"', company_input)
        if len(companies) == 0:
            print("Invalid input. Please enter at least one company name.")
        else:
            return [companies]
def get_doc_path():
    while True:
        doc_path = input("Enter the path to the document (e.g., C:\\Users\\Mudassir\\Desktop\\New folder\\New Sample.docx): ")
        if doc_path:
            return doc_path
        else:
            print("Invalid input. Please enter a valid document path.")
def main():
    doc_path = get_doc_path()
    print(doc_path)
    if not os.path.exists(doc_path):
        print("Document path is invalid.")
        exit()
    doc = None
    try:
        doc = Document(doc_path)
    except Exception as e:
        print(f"Error loading the document: {e}")
        exit()
    transformer_market = get_segments()
    print(transformer_market)
    company = get_companies()
    print(company)
    # Remove hyperlinks from the document
    for paragraph in doc.paragraphs:
        remove_hyperlinks(paragraph)
    replacetext(transformer_market,company, doc)
    delete_segments_and_blank_pages(doc)

    new_doc_path = os.path.join(os.curdir, 'Output.docx')
    doc.save(new_doc_path)
    print(f"Modifications saved successfully to {new_doc_path}.")

main()